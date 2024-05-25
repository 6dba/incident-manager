"""
Сервис файловой обработки инцидентов
"""
__author__ = "6dba"
__date__ = "07/05/2024"

import os
from pathlib import Path
from typing import Optional
from docx import Document
from docx.shared import Pt

from core.messaging.broker import Exchanges, Queues
from core.repositories.base.incident import IncidentModel
from core.service import BaseService
from services.manager.client import FTPClient, SMBClient
from services.manager.templates.messages import TO_REPLACE


class ManagerService(BaseService):
    """
    Сервис файлов
    """
    def __init__(self):
        super().__init__()
        # Обменник, в который сервис публикует события
        self.exchange = Exchanges.MANAGER.value
        self.queue = Queues.MANAGER.value
        self.__service_dir_from_root = 'services/manager'
        self.__template_dir = f'{self.__service_dir_from_root}/templates/'
        # {Название шаблона: Путь к файлу шаблона}
        self.templates = {
            os.path.basename(p).removesuffix('.docx'): self.__template_dir + p
            for p in os.listdir(self.__template_dir) if p.endswith('.docx')
        }
        self.register_handlers()

    def register_handlers(self):
        """
        Регистрация обработчиков сервиса
        """

        @self.broker.handle(self.queue, Exchanges.PROCESSOR.value)  # Подписка события сервиса пост-обработки данных
        async def manage(incidents: list[IncidentModel]):
            await self.__manage(incidents)

    async def __manage(self, incidents: list[IncidentModel]):
        """
        Обработка инцидентов, заполнение шаблона

        :param list[IncidentModel] incidents:
        """
        processed_incidents = []
        processed_docs = {}

        for incident in incidents:
            doc = self.__resolve_document(incident)
            if not doc:
                continue
            table = doc.tables[0]
            # Заполнение таблицы документа, по данным инцидента
            self.__process_table(incident, table)
            processed_incidents.append(incident)
            processed_docs.update({f'{incident.id}_{incident.gossopka_incident_type}.docx': doc})

        if processed_incidents:
            # Если появились обработанные документы - публикуем связанные с ними инциденты,
            # как способ сказать о том, какие инциденты были обработаны менеджером
            await self.broker.publish(exchange=self.exchange, message=processed_incidents)

        if processed_docs:
            # Документы публикуем на указанный FTP сервер, для потребителей
            self.__to_file_resources(processed_docs)

    def __to_file_resources(self, documents: dict[str: Document]):
        """
        Передача обработанного файла на файловый ресурс
        :param documents:
        :return:
        """
        for title, document in documents.items():
            temp_path = f'{self.__service_dir_from_root}/temp'
            Path(temp_path).mkdir(parents=True, exist_ok=True)
            local_path = f'{temp_path}/{title}'
            # Временно сохраняем документ
            document.save(local_path)
            try:
                _, _ = FTPClient().upload(local_path, title), SMBClient().upload(local_path, title)
            except Exception as e:
                self.logger.exception(e)
            finally:
                # Удаление временного файла
                os.remove(local_path)

    @staticmethod
    def __process_table(incident: IncidentModel, table: Document):
        """
        Заполнение таблицы документа, по данным инцидента

        :param IncidentModel incident:
        :param table:
        :return:
        """
        replaced_titles = []

        for row in table.rows:
            # Колонка с названием строки
            title = row.cells[0].text

            if title in replaced_titles:
                # Особенности строки с двумя колонками, когда вторая колонка содержит несколько строк,
                # которые по логике таблицы относятся к одной строке
                # _________________________________________________________________________________
                #                                         | Да
                # Необходимость привлечения сил ГосСОПКА  -----------------------------------------
                #                                         | Нет
                # _________________________________________________________________________________
                # Итерация по ячейкам в таком случае происходит попарно, то есть фактически строк будет две:
                # ('Необходимость привлечения сил ГосСОПКА', 'Да'), ('Необходимость привлечения сил ГосСОПКА', 'Нет')
                # Соответственно при первой итерации, выбираем нужно значение,
                # а при второй итерации - удаляем уже обработанную строку
                table._tbl.remove(row._tr)
                continue

            replaced = TO_REPLACE.get(title, lambda _: None)(incident)
            if not replaced:
                # Если для данной строки не нужно производить замен
                continue

            # Замена текста в ячейке таблицы
            row.cells[1].text = replaced
            # Изменение размера шрифта
            row.cells[1].paragraphs[0].runs[0].font.size = Pt(12)
            replaced_titles.append(title)

    def __resolve_document(self, incident: IncidentModel) -> Optional[Document]:
        """
        Разрешение документа-шаблона для инцидента
        :param IncidentModel incident:
        :return: Объект документа .docx
        :rtype: Document
        :raise: ValueError
        """
        if incident.gossopka_incident_type.lower() not in list(map(lambda t: t.lower(), self.templates.keys())):
            # Если тип инцидента не соответствует ни одному шаблону
            self.logger.warning(
                f'Не найдено ни одного подходящего шаблона для типа инцидента {incident.gossopka_incident_type}'
                f'Перечень шаблонов: {self.templates}'
            )
            return None

        # Создаем объект документа для выбранного шаблона
        return Document(self.templates.get(incident.gossopka_incident_type, ''))

    async def start(self):
        """
        Рабочий цикл сервиса
        """
        await self.app.run()
